
from genericpath import exists
from pydoc import doc

import pandas as pd
from docx.shared import Cm, Inches
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph
from docx.enum.text import WD_ALIGN_PARAGRAPH

"""
Este módulo contiene funciones para reemplazar marcadores de posición en párrafos de Word, incluyendo la inserción de imágenes con o sin títulos.
Conceptos clave:
paragraph - Un párrafo en Word, que puede contener múltiples runs.
run - Es un fragmento de texto dentro de un párrafo que comparte el mismo formato.
Un párrafo puede tener múltiples runs con diferentes formatos.
ejemplo: "El valor es <<valor>>" podría estar dividido en 3 runs: ["El valor es ", "<<valor>>", ""]

"""

# def insertar_parrafo_despues(paragraph, texto="", centrado=False):
#     """Inserta un nuevo párrafo después del párrafo dado usando manipulación XML.
    
#     Args:
#         paragraph: El párrafo después del cual insertar.
#         texto: El texto del nuevo párrafo (opcional).
#         centrado: Si True, centra el texto del párrafo.
    
#     Returns:
#         El elemento XML del nuevo párrafo creado.
#     """
#     # Obtener el elemento del párrafo actual
#     p_element = paragraph._element
#     # Obtener el elemento padre
#     parent = p_element.getparent()
#     # Crear un nuevo elemento de párrafo
#     nuevo_p = OxmlElement('w:p')
    
#     # Si necesita estar centrado
#     if centrado:
#         pPr = OxmlElement('w:pPr')
#         jc = OxmlElement('w:jc')
#         jc.set(qn('w:val'), 'center')
#         pPr.append(jc)
#         nuevo_p.append(pPr)
    
#     # Si hay texto, agregarlo
#     if texto:
#         run = OxmlElement('w:r')
#         text_elem = OxmlElement('w:t')
#         text_elem.text = texto
#         run.append(text_elem)
#         nuevo_p.append(run)
    
#     # Insertar el nuevo párrafo después del actual
#     parent.insert(parent.index(p_element) + 1, nuevo_p)
    
#     return nuevo_p


def aux_insertar_figura_sin_titulo(paragraph, key, ruta_imagen, ancho_inches=6):
    """Inserta UNA imagen en un párrafo de Word SIN título (caso de array con un solo elemento).
    
    Args:
        paragraph: El párrafo del documento Word donde se buscará el marcador.
        key: El marcador de posición a buscar (por ejemplo, "<<fig_ejecucion_campania>>").
        ruta_imagen: La ruta del archivo de imagen a insertar.
        ancho_cm: El ancho de la imagen en centímetros (por defecto 15).
    
    Returns:
        True si se insertó la imagen, False si no se encontró el marcador.
    """
    full_text = "".join(run.text for run in paragraph.runs)
    
    if key not in full_text:
        return False
    
    # Limpiar el párrafo (borrar todos los runs)
    for run in paragraph.runs:
        run.text = ""
    
    # Insertar la imagen en el párrafo
    run = paragraph.runs[0] if paragraph.runs else paragraph.add_run()
    run.add_picture(ruta_imagen, width=Inches(ancho_inches))
    
    # Insertar un salto de línea (párrafo vacío) después
    # insertar_parrafo_despues(paragraph, texto="")
    
    return True

def aux_reemplazar_variable_en_parrafo(paragraph, key, value):
    """Reemplaza un marcador de posición en un párrafo de Word PRESERVANDO el formato.
    
    Args:
        paragraph: El párrafo donde buscar el marcador.
        key: El marcador de posición a buscar (por ejemplo, "<<orden_de_servicio>>").
        value: El texto que lo reemplazará (i.e., 100).
    
    Esta función preserva el formato del run donde EMPIEZA el marcador.
    Maneja tanto marcadores dentro de un solo run como divididos entre múltiples runs.
    """
    # Unir todo el texto del párrafo para verificar si este tiene al marcador
    full_text = "".join(run.text for run in paragraph.runs)
    if key not in full_text:
        return paragraph
    
    # Convertir value a string
    new_value = str(value[0]) if isinstance(value, list) else str(value)
    
    # CASO 1: Intentar el caso simple primero (todo en un run)
    for irun, run in enumerate(paragraph.runs):
        if key in run.text:
            run.text = run.text.replace(key, new_value)
        
        elif "<<" in run.text: # Si el marcador está dividido
            variable_word = "".join([paragraph.runs[irun].text, paragraph.runs[irun+1].text, paragraph.runs[irun+2].text])
            if key in variable_word:
                paragraph.runs[irun].text = variable_word.replace(key, new_value)
                paragraph.runs[irun+1].text = ""
                paragraph.runs[irun+2].text = ""
                return paragraph

    
def aplicar_formato_celda(celda, formato, texto):
   return


def rellenar_tabla(doc, nombre_marcador, diccionario_de_reemplazos):
    """Rellena una tabla dinámicamente buscando un marcador e insertando datos de un dataframe.
    
    Args:
        doc: El documento de Word (objeto Document).
        nombre_marcador: El nombre del marcador a buscar en la tabla (ej: "<<tabla1>>").
        diccionario_de_reemplazos: Diccionario con los datos a insertar en la tabla.
    
    Descripción:
        La función busca el marcador en todas las tablas del documento.
        Una vez encontrado, guarda el formato de la fila 1 (que contiene el marcador),
        elimina esa fila, y luego inserta todas las filas del dataframe en la tabla
        aplicando el formato guardado.
        
    Ejemplo:
        df = pd.DataFrame({
            'secuencia': [0, 1, 2],
            'localizacion': ['BOT-01', 'BOT-02', 'BOT-03'],
            'lat_plan': [18.5, 18.6, 18.7]
        })
        rellenar_tabla(doc, "<<tabla1>>", df)
        # Resultado: La tabla tendrá 4 filas (1 de encabezado + 3 de datos)
    """
    df = pd.DataFrame(diccionario_de_reemplazos)
    cantidad_de_filas = len(df)
    
    # Buscar la tabla que contiene el marcador
    for table in doc.tables:
        marcador = table.rows[0].cells[0].text
        if nombre_marcador in marcador:
            for idx, row in df.iterrows():
                for icol, col in enumerate(df.columns): #Agrego los datos del dataframe a la tabla
                    table.rows[idx+1].cells[icol].text = str(row[col])
                    table.rows[idx+1].cells[icol].paragraphs[0].style = doc.styles['texto_tablas_centrado'] # Agregar el estilo daña el alineado vertical y el tamaño de la fila
                    table.rows[idx+1].cells[icol].vertical_alignment = WD_ALIGN_PARAGRAPH.CENTER 
                    table.rows[idx+1].height = 288290
                    
                    if nombre_marcador == "<<tabla2>>" and icol == 1:
                        table.rows[idx+1].cells[icol].paragraphs[0].style = doc.styles['texto_tablas_justificado'] # Agregar el estilo daña el alineado vertical y el tamaño de la fila
                        table.rows[idx+1].cells[icol].vertical_alignment = WD_ALIGN_PARAGRAPH.CENTER
                        table.rows[idx+1].height = 288290 # Ajustar la altura de la fila según sea necesario
                        
                if idx+1 < cantidad_de_filas: # Agrego una fila vacía para la siguiente iteración (si es que hay más filas por agregar)    
                    table.add_row().cells # Al terminar agrego una fila vacía para la siguiente iteración (si es que hay más filas por agregar)

            # PASO 2: Eliminar la última fila
            # table._element.remove(table.rows[-1]._element)
            # 
            
            break
        
    
    return doc


def reemplazar_en_word(doc, diccionario_de_reemplazos):
    """Reemplaza los marcadores de posición en un documento de Word utilizando un diccionario de reemplazos.
    doc es el documento de Word (objeto Document).
    diccionario_de_reemplazos es un diccionario donde las claves son los marcadores de posición a buscar
    (por ejemplo, "<<orden_de_servicio>>") y los valores son los textos que los reemplazarán (i.e., 100).
    
    Para las figuras, el valor debe ser una lista de diccionarios:
    - Lista con un solo elemento: inserta la figura SIN título
    - Lista con varios elementos: inserta las figuras CON sus títulos
    
    Cada diccionario debe tener las keys: "ruta", "titulo", "tamanio"
    """
    # Para cada párrafo en el documento, reemplaza los marcadores de posición utilizando el diccionario
    for variable, dato in diccionario_de_reemplazos.items():
        for parrafo in doc.paragraphs:
            if variable in parrafo.text:
                
                if "fig" not in variable: # Si el marcador no es de figura, reemplazo normal
                    aux_reemplazar_variable_en_parrafo(parrafo, variable, dato)
                
                if "fig" in variable: # Si el marcador es de figura
                    aux_insertar_figura_sin_titulo(parrafo, variable, dato, 4)
                #         else: # Varias figuras CON títulos
                #             aux_insertar_figuras_con_titulo(parrafo, variable, dato)
                # else: # No es un marcador de figura, reemplazo normal
                #     aux_reemplazar_en_parrafo(parrafo, variable, dato)
    
    # Ahora busco el marcador tabla1 y ejecuto la función que reemplaza el texto y agrega las filas necesarias para la tabla1
    
    # Ahora busco el marcador tabla2 y ejecuto la función que reemplaza el texto y agrega las filas necesarias para la tabla2
    
    

    
    return doc